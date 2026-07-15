#!/usr/bin/env python3
"""Minimal, self-contained Markdown -> .docx renderer tuned for this report.
Handles: # title, ##/###/#### headings, GFM tables (with a |---| separator row),
- bullets, **bold**/*italic* inline, --- rules, and plain paragraphs.
Usage: python3 md2docx.py input.md output.docx"""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

IN, OUT = sys.argv[1], sys.argv[2]
with open(IN, encoding="utf-8") as f:
    lines = f.read().split("\n")

doc = Document()
# base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
for lvl, sz, col in [("Heading 1", 15, (0x1F, 0x3A, 0x5F)),
                     ("Heading 2", 12.5, (0x1F, 0x3A, 0x5F)),
                     ("Heading 3", 11, (0x2E, 0x4A, 0x6E))]:
    st = doc.styles[lvl]
    st.font.size = Pt(sz); st.font.bold = True
    st.font.color.rgb = RGBColor(*col); st.font.name = "Calibri"

INLINE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")

def add_runs(paragraph, text):
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith("**"):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        else:
            r = paragraph.add_run(tok[1:-1]); r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])

def is_table_row(s):
    return s.strip().startswith("|") and s.strip().endswith("|")

def is_sep_row(s):
    body = s.replace("|", "").replace("-", "").replace(":", "").strip()
    return body == "" and "-" in s

def cells(s):
    parts = [c.strip() for c in s.strip().split("|")]
    return parts[1:-1]  # drop the empty ends

def render_table(block):
    header = cells(block[0])
    body = [cells(r) for r in block[2:]]  # skip header + separator
    ncol = len(header)
    t = doc.add_table(rows=1, cols=ncol)
    try:
        t.style = "Light Grid Accent 1"
    except Exception:
        t.style = "Table Grid"
    t.autofit = True
    hdr = t.rows[0].cells
    for j, c in enumerate(header):
        p = hdr[j].paragraphs[0]
        add_runs(p, c)
        for run in p.runs:
            run.bold = True
    for row in body:
        rc = t.add_row().cells
        for j in range(ncol):
            val = row[j] if j < len(row) else ""
            add_runs(rc[j].paragraphs[0], val)
    doc.add_paragraph()  # spacer after table

i = 0
n = len(lines)
while i < n:
    line = lines[i]
    s = line.strip()
    if is_table_row(line):
        block = []
        while i < n and is_table_row(lines[i]):
            block.append(lines[i]); i += 1
        # require a separator row as the 2nd line to treat as table
        if len(block) >= 2 and is_sep_row(block[1]):
            render_table(block)
        else:
            for b in block:
                add_runs(doc.add_paragraph(), b)
        continue
    if s == "":
        i += 1; continue
    if s == "---":
        doc.add_paragraph().add_run("").add_break  # light spacer
        i += 1; continue
    if s.startswith("#### "):
        doc.add_heading(s[5:], level=3); i += 1; continue
    if s.startswith("### "):
        h = doc.add_heading("", level=3); add_runs(h, s[4:]); i += 1; continue
    if s.startswith("## "):
        h = doc.add_heading("", level=2); add_runs(h, s[3:]); i += 1; continue
    if s.startswith("# "):
        h = doc.add_heading("", level=1); add_runs(h, s[2:])
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        i += 1; continue
    if s.startswith("> "):
        p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(s[2:]); r.italic = True; i += 1; continue
    if s.startswith("- "):
        p = doc.add_paragraph(style="List Bullet"); add_runs(p, s[2:]); i += 1; continue
    # normal paragraph
    add_runs(doc.add_paragraph(), s)
    i += 1

doc.save(OUT)
print("WROTE", OUT)
# quick stats
print("paragraphs:", len(doc.paragraphs), "tables:", len(doc.tables))
